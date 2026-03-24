#!/usr/bin/env python3
"""Generate AI Implementation Guide for Engineers - Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(12)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

def add_checkbox(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' {text}')
    else:
        p.add_run(f'☐ {text}')
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f'💡 {text}')
    run.italic = True
    run.font.color.rgb = RGBColor(0x2e, 0x7d, 0x32)
    return p

def add_eng_example(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f'⚙️ Применение в проектировании: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
    p.add_run(text)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    return table


# ══════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Полное руководство по внедрению ИИ\nв инженерную работу и проектирование')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('От первого запуска до полноценной интеграции\nв рабочий процесс инженера')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('Версия 1.0 — Март 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════
# TABLE OF CONTENTS placeholder
# ══════════════════════════════════════════════
doc.add_heading('Содержание', level=1)
toc_items = [
    'Часть 1. Облачный ИИ-ассистент (OpenClaw)',
    '  Этап 0. Подготовка инфраструктуры',
    '  Этап 1. Установка и первый запуск',
    '  Этап 2. Настройка под инженерные задачи',
    '  Этап 3. Продвинутое использование',
    '  Этап 4. Интеграция в проектный процесс',
    'Часть 2. Локальная ИИ-модель (без облака)',
    '  Этап 0. Оценка оборудования',
    '  Этап 1. Установка Ollama',
    '  Этап 2. Интерфейсы для работы',
    '  Этап 3. Оптимизация производительности',
    '  Этап 4. Продвинутые сценарии',
    'Часть 3. ИИ в инженерном проектировании',
    '  Сценарии применения по дисциплинам',
    '  Примеры промптов для инженерных задач',
    '  Сравнение моделей для инженерных задач',
    'Приложения',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════
# ВВЕДЕНИЕ
# ══════════════════════════════════════════════
doc.add_heading('Введение', level=1)
doc.add_paragraph(
    'Этот документ — пошаговое руководство для инженера, который хочет внедрить '
    'искусственный интеллект в свою работу. Он покрывает два направления:'
)
doc.add_paragraph('• Облачные ИИ-модели — мощные, не требуют мощного железа, работают через API')
doc.add_paragraph('• Локальные ИИ-модели — полная конфиденциальность, работают без интернета, бесплатно')
doc.add_paragraph()
doc.add_paragraph(
    'Особый акцент сделан на применение в инженерном проектировании: '
    'SCADA/АСУ ТП, электротехника, КИПиА, промышленная автоматизация, '
    'подготовка технической документации, расчёты и анализ данных.'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# ЧАСТЬ 1
# ══════════════════════════════════════════════
doc.add_heading('Часть 1. Облачный ИИ-ассистент (OpenClaw)', level=1)
doc.add_paragraph(
    'OpenClaw — это open-source платформа, которая превращает ИИ-модель в полноценного '
    'персонального ассистента. Вы подключаете его к Telegram, Discord или Signal, и '
    'он становится вашим помощником с доступом к файлам, командной строке, интернету и памятью.'
)

# ── Этап 0 ──
doc.add_heading('Этап 0. Подготовка инфраструктуры', level=2)

doc.add_heading('Выбор сервера', level=3)
doc.add_paragraph(
    'OpenClaw — это «мозг» вашего ассистента. Он работает на сервере 24/7 и пересылает '
    'ваши сообщения ИИ-модели. Сам по себе он лёгкий (Node.js процесс), основная '
    'вычислительная нагрузка лежит на облачном провайдере.'
)
add_table(doc,
    ['Вариант', 'Стоимость', 'Плюсы', 'Минусы'],
    [
        ['VPS (Timeweb, Selectel, Hetzner)', '300–800 ₽/мес', 'Работает 24/7, статический IP', 'Нужна базовая настройка'],
        ['Домашний ПК / Raspberry Pi', 'Бесплатно', 'Полный контроль', 'Нужен постоянный интернет'],
        ['Старый ноутбук', 'Бесплатно', 'Уже есть', 'Шумит, потребляет энергию'],
    ]
)
doc.add_paragraph()
doc.add_paragraph('Минимальные требования к серверу:')
doc.add_paragraph('• ОС: Ubuntu 22.04+ или Debian 12+')
doc.add_paragraph('• RAM: 2 ГБ (рекомендовано 4 ГБ)')
doc.add_paragraph('• Диск: 20 ГБ SSD')
doc.add_paragraph('• Сеть: постоянное интернет-соединение')
doc.add_paragraph('• CPU: 1–2 ядра (OpenClaw не нагружает процессор)')

add_eng_example(doc,
    'Для инженерного отдела из 3–5 человек достаточно одного VPS за 500 ₽/мес. '
    'Каждый инженер подключается через свой Telegram, а OpenClaw маршрутизирует запросы.'
)

doc.add_heading('Выбор провайдера ИИ-модели', level=3)
doc.add_paragraph(
    'OpenClaw не содержит собственную модель — он подключается к облачным провайдерам. '
    'Это значит, что вы выбираете «мозг» отдельно от «тела».'
)
add_table(doc,
    ['Провайдер', 'Лучшие модели', 'Стоимость', 'Для инженерии'],
    [
        ['Anthropic', 'Claude Sonnet 4, Claude Opus 4', '$3–15 / 1M токенов', 'Лучший для текстов и анализа (оценка автора)'],
        ['OpenAI', 'GPT-4.1, o3, Codex', '$2–10 / 1M токенов', 'Хорош для кода и расчётов (оценка автора)'],
        ['Google', 'Gemini 2.5 Pro/Flash', '$1–7 / 1M токенов', 'Большое контекстное окно (до 1M токенов)'],
        ['OpenRouter', 'Все + бесплатные (Qwen3, DeepSeek)', 'Наценка ~10%', 'Один ключ на всё, есть бесплатные модели'],
    ]
)
doc.add_paragraph()
add_note(doc,
    'Что такое токены? Это единицы текста для ИИ. 1000 токенов ≈ 750 слов на английском '
    'или ≈ 500 слов на русском. При активном ежедневном использовании расход: 100–500 тыс. '
    'токенов в день = $5–20 в месяц.'
)

doc.add_heading('Необходимые навыки', level=3)
doc.add_paragraph('Для базовой установки достаточно:')
doc.add_paragraph('• Умение подключиться к серверу по SSH (PuTTY, Terminal)')
doc.add_paragraph('• Базовые команды: cd, ls, nano (или vim)')
doc.add_paragraph('• Умение скопировать и вставить команду в терминал')
doc.add_paragraph('• Аккаунт Telegram (для создания бота через @BotFather)')

doc.add_page_break()

# ── Этап 1 ──
doc.add_heading('Этап 1. Установка и первый запуск', level=2)

doc.add_heading('Шаг 1. Установка Node.js', level=3)
doc.add_paragraph('OpenClaw работает на Node.js. Установите версию 20 или новее:')
add_code(doc, '# Для Ubuntu/Debian:')
add_code(doc, 'curl -fsSL https://deb.nodesource.com/setup_22.x | sudo bash -')
add_code(doc, 'sudo apt install -y nodejs')
add_code(doc, '')
add_code(doc, '# Проверка:')
add_code(doc, 'node --version   # должно быть v22.x.x')
add_code(doc, 'npm --version    # должно быть 10.x.x')

doc.add_heading('Шаг 2. Установка OpenClaw', level=3)
add_code(doc, 'sudo npm install -g openclaw')
add_code(doc, '')
add_code(doc, '# Проверка:')
add_code(doc, 'openclaw --version')

doc.add_heading('Шаг 3. Мастер настройки', level=3)
doc.add_paragraph(
    'Запустите интерактивный мастер, который проведёт через все настройки:'
)
add_code(doc, 'openclaw setup')
doc.add_paragraph('Мастер попросит:')
doc.add_paragraph('1. Выбрать провайдера модели (Anthropic, OpenAI, Google, OpenRouter)')
doc.add_paragraph('2. Ввести API-ключ (получить на сайте провайдера)')
doc.add_paragraph('3. Выбрать модель по умолчанию')
doc.add_paragraph('4. Подключить мессенджер (Telegram, Discord, Signal)')

doc.add_heading('Шаг 4. Создание Telegram-бота', level=3)
doc.add_paragraph('Если выбран Telegram:')
doc.add_paragraph('1. Откройте Telegram, найдите @BotFather')
doc.add_paragraph('2. Отправьте /newbot')
doc.add_paragraph('3. Введите имя бота (например: «Мой Инженерный Ассистент»)')
doc.add_paragraph('4. Введите username (например: my_engineer_bot)')
doc.add_paragraph('5. Скопируйте полученный токен')
doc.add_paragraph('6. Вставьте токен в мастер настройки OpenClaw')

doc.add_heading('Шаг 5. Первый запуск', level=3)
add_code(doc, '# Запуск в фоновом режиме:')
add_code(doc, 'openclaw gateway start')
add_code(doc, '')
add_code(doc, '# Проверка статуса:')
add_code(doc, 'openclaw status')
add_code(doc, '')
add_code(doc, '# Просмотр логов:')
add_code(doc, 'openclaw gateway logs')
doc.add_paragraph()
doc.add_paragraph(
    'Откройте Telegram, найдите вашего бота и напишите «Привет». '
    'Если всё настроено правильно — бот ответит.'
)

add_eng_example(doc,
    'Первое, что стоит попросить бота: «Объясни мне, что ты умеешь для инженера-проектировщика». '
    'Модель расскажет о своих возможностях в контексте вашей работы.'
)

doc.add_page_break()

# ── Этап 2 ──
doc.add_heading('Этап 2. Настройка под инженерные задачи', level=2)

doc.add_heading('Настройка «души» ассистента (SOUL.md)', level=3)
doc.add_paragraph(
    'Файл SOUL.md определяет характер и стиль общения. Для инженерного ассистента '
    'рекомендуется задать технический, точный стиль.'
)
doc.add_paragraph('Пример содержимого:')
add_code(doc, '# SOUL.md')
add_code(doc, '')
add_code(doc, '## Кто я')
add_code(doc, 'Я — инженерный ассистент. Специализация: промышленная автоматизация,')
add_code(doc, 'проектирование АСУ ТП, электротехника, КИПиА.')
add_code(doc, '')
add_code(doc, '## Стиль общения')
add_code(doc, '- Точный и технически грамотный')
add_code(doc, '- Использую правильную терминологию (ГОСТ, МЭК/IEC)')
add_code(doc, '- Даю ссылки на нормативные документы')
add_code(doc, '- Предупреждаю о рисках и ограничениях')
add_code(doc, '- Не упрощаю, если спрашивает специалист')

doc.add_heading('Профиль пользователя (USER.md)', level=3)
doc.add_paragraph('Расскажите ассистенту о себе — это улучшит качество ответов:')
add_code(doc, '# USER.md')
add_code(doc, '')
add_code(doc, '- Имя: [Ваше имя]')
add_code(doc, '- Должность: Инженер-проектировщик АСУ ТП')
add_code(doc, '- Компания: [Название]')
add_code(doc, '- Специализация: SCADA, ПЛК Siemens/Schneider, КИПиА')
add_code(doc, '- Инструменты: EPLAN, AutoCAD, TIA Portal, Step 7')
add_code(doc, '- Стандарты: ГОСТ 21.408, ГОСТ 34, IEC 61131-3, IEC 62443')

doc.add_heading('Рабочее правило поведения (AGENTS.md)', level=3)
doc.add_paragraph('Добавьте в AGENTS.md инженерные правила:')
doc.add_paragraph('• Всегда указывать источник (ГОСТ, СП, нормативный документ)')
doc.add_paragraph('• При расчётах показывать формулы и единицы измерения')
doc.add_paragraph('• Предупреждать, если данных недостаточно для однозначного ответа')
doc.add_paragraph('• Не давать рекомендаций по безопасности без оговорок')

doc.add_heading('База знаний проекта', level=3)
doc.add_paragraph(
    'Загрузите рабочие документы в workspace ассистента — это его «справочная библиотека»:'
)
doc.add_paragraph('• Технические задания (ТЗ) текущих проектов')
doc.add_paragraph('• Шаблоны документов (пояснительные записки, спецификации)')
doc.add_paragraph('• Каталоги оборудования (PDF, Excel)')
doc.add_paragraph('• Внутренние стандарты предприятия')
doc.add_paragraph('• Типовые решения и наработки')
add_code(doc, '# Структура workspace для инженера:')
add_code(doc, 'docs/')
add_code(doc, '  standards/       # ГОСТы, СП, нормативы')
add_code(doc, '  catalogs/        # Каталоги оборудования')
add_code(doc, '  templates/       # Шаблоны документов')
add_code(doc, 'projects/')
add_code(doc, '  current_project/ # Файлы текущего проекта')
add_code(doc, '  archive/         # Завершённые проекты')

add_eng_example(doc,
    'Загрузите каталог датчиков Endress+Hauser в docs/catalogs/. Теперь можно спросить: '
    '«Подбери датчик давления для пара 10 бар, 180°C, присоединение G½» — и ассистент '
    'найдёт подходящий прибор прямо из каталога.'
)

doc.add_page_break()

# ── Этап 3 ──
doc.add_heading('Этап 3. Продвинутое использование', level=2)

doc.add_heading('Навыки (Skills) через ClawHub', level=3)
doc.add_paragraph(
    'Skills — это плагины, расширяющие возможности ассистента. '
    'Устанавливаются одной командой:'
)
add_code(doc, '# Поиск навыков:')
add_code(doc, 'clawhub search "github"')
add_code(doc, '')
add_code(doc, '# Установка:')
add_code(doc, 'clawhub install github')
doc.add_paragraph('Полезные навыки для инженера:')
doc.add_paragraph('• github — работа с Git-репозиториями (код, документация)')
doc.add_paragraph('• weather — прогноз погоды (для полевых работ и пусконаладки)')
doc.add_paragraph('• presentation-designer — создание презентаций PowerPoint')

doc.add_heading('Автоматизация через Cron', level=3)
doc.add_paragraph('Настройте регулярные задачи:')
doc.add_paragraph('• Утренний дайджест — сводка по проектам, письмам, задачам')
doc.add_paragraph('• Мониторинг серверов — проверка SCADA-серверов (если есть SSH-доступ)')
doc.add_paragraph('• Напоминания — дедлайны, согласования, поверки приборов')

add_eng_example(doc,
    'Настройте ежедневный cron: «Проверь сроки поверки приборов в файле '
    'projects/calibration_schedule.xlsx и напомни, если что-то истекает в ближайшие 30 дней». '
    'Ассистент будет каждое утро присылать отчёт в Telegram.'
)

doc.add_heading('Работа с GitHub', level=3)
doc.add_paragraph(
    'Если ваша проектная документация или код хранится в Git:'
)
doc.add_paragraph('• Ассистент может коммитить изменения')
doc.add_paragraph('• Создавать и ревьюить Pull Request-ы')
doc.add_paragraph('• Отслеживать Issues')
doc.add_paragraph('• Автоматически проверять CI/CD пайплайны')

doc.add_heading('Sub-agents (параллельные агенты)', level=3)
doc.add_paragraph(
    'Для сложных задач OpenClaw может запускать несколько ИИ-агентов параллельно:'
)
doc.add_paragraph('• Один агент пишет пояснительную записку')
doc.add_paragraph('• Другой составляет спецификацию оборудования')
doc.add_paragraph('• Третий проверяет расчёты')
doc.add_paragraph('Результаты собираются в один ответ.')

doc.add_page_break()

# ── Этап 4 ──
doc.add_heading('Этап 4. Интеграция в проектный процесс', level=2)

doc.add_heading('Определите задачи для ИИ', level=3)
doc.add_paragraph('Начните с задач, где ИИ даёт максимальную отдачу:')

add_table(doc,
    ['Задача', 'Экономия времени*', 'Сложность внедрения'],
    [
        ['Написание пояснительных записок', '60–80%*', 'Низкая'],
        ['Подбор оборудования по каталогам', '40–60%*', 'Средняя'],
        ['Проверка документации на ошибки', '50–70%*', 'Низкая'],
        ['Генерация кода ПЛК (ST/SCL)', '30–50%*', 'Средняя'],
        ['Расчёты (кабельные, гидравлические)', '40–60%*', 'Средняя'],
        ['Перевод технической документации', '70–90%*', 'Низкая'],
        ['Составление ТКП / КП', '50–70%*', 'Средняя'],
        ['Анализ P&ID и функциональных схем', '20–40%*', 'Высокая'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('* Проценты экономии — экспертная оценка автора. Реальная экономия зависит от сложности задачи, качества промптов и опыта работы с ИИ.')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_heading('Настройте шаблоны и процедуры', level=3)
doc.add_paragraph(
    'Создайте в workspace шаблоны для типовых задач. Ассистент будет использовать '
    'их как основу:'
)
doc.add_paragraph('• Шаблон пояснительной записки по ГОСТ 34')
doc.add_paragraph('• Шаблон спецификации оборудования')
doc.add_paragraph('• Шаблон акта испытаний')
doc.add_paragraph('• Чек-лист приёмки SCADA-системы')

doc.add_heading('Наладьте обратную связь', level=3)
doc.add_paragraph(
    'ИИ обучается в рамках сессии на ваших правках. Чем больше вы его поправляете, '
    'тем точнее он становится в рамках проекта:'
)
doc.add_paragraph('• Поправляйте терминологию: «Не "сенсор", а "датчик"»')
doc.add_paragraph('• Указывайте на ошибки: «Здесь неправильная маркировка по ГОСТ 21.408»')
doc.add_paragraph('• Хвалите хорошие результаты — модель учитывает это')
doc.add_paragraph('• Сохраняйте удачные промпты в docs/prompts/')

doc.add_page_break()

# ══════════════════════════════════════════════
# ЧАСТЬ 2
# ══════════════════════════════════════════════
doc.add_heading('Часть 2. Локальная ИИ-модель (без облака)', level=1)
doc.add_paragraph(
    'Локальная модель работает полностью на вашем компьютере. Никакие данные не покидают '
    'вашу сеть. Это критично для работы с конфиденциальной проектной документацией, '
    'чертежами заказчика и коммерческими тайнами.'
)

# ── Этап 0 ──
doc.add_heading('Этап 0. Оценка оборудования', level=2)

doc.add_heading('Что определяет производительность', level=3)
doc.add_paragraph(
    'Локальная модель — это нейросеть, которая работает на вашем железе. '
    'Главные факторы:'
)
doc.add_paragraph('• GPU (видеокарта) — основной ускоритель. Чем больше VRAM, тем крупнее модель')
doc.add_paragraph('• RAM (оперативная память) — если нет GPU, модель работает в RAM на CPU')
doc.add_paragraph('• CPU — влияет на скорость, если нет GPU')
doc.add_paragraph('• SSD — быстрая загрузка модели с диска')

doc.add_heading('Таблица совместимости', level=3)
add_table(doc,
    ['Конфигурация', 'Подходящие модели', 'Скорость', 'Рекомендация'],
    [
        ['8 ГБ RAM, без GPU', 'Qwen 3.5 4B', '3–8 tok/s', 'Только для простых задач'],
        ['16 ГБ RAM, без GPU', 'Qwen 3.5 9B (q4)', '5–12 tok/s', 'Базовое использование'],
        ['16 ГБ RAM + GPU 8 ГБ', 'Qwen 3.5 35B-A3B (MoE)', '20–40 tok/s', 'Комфортная работа, MoE эффективен'],
        ['32 ГБ RAM + GPU 12 ГБ', 'Qwen 3.5 27B (dense)', '15–30 tok/s', 'Хорошее качество'],
        ['64 ГБ RAM + GPU 24 ГБ', 'Qwen 3.5 122B-A10B, DeepSeek-R1 32B', '10–25 tok/s', 'Отличное качество'],
        ['GPU 48 ГБ+ / 2×24 ГБ', 'Qwen 3.5 397B-A17B', '5–15 tok/s', 'Близко к облачным'],
    ]
)
doc.add_paragraph()
add_note(doc,
    'tok/s (tokens per second) — скорость генерации. 15+ tok/s — комфортно для диалога. '
    '5–10 — терпимо. Меньше 5 — лучше оставлять задачу и возвращаться за результатом.'
)

doc.add_heading('Проверка вашего оборудования', level=3)
add_code(doc, '# NVIDIA GPU:')
add_code(doc, 'nvidia-smi')
add_code(doc, '')
add_code(doc, '# AMD GPU:')
add_code(doc, 'rocm-smi')
add_code(doc, '')
add_code(doc, '# RAM:')
add_code(doc, 'free -h')
add_code(doc, '')
add_code(doc, '# CPU:')
add_code(doc, 'lscpu | grep "Model name"')
add_code(doc, '')
add_code(doc, '# На Windows:')
add_code(doc, '# Диспетчер задач → Производительность → GPU / Память')

add_eng_example(doc,
    'Типичная инженерная рабочая станция (i7/Ryzen 7, 32 ГБ RAM, RTX 3060 12 ГБ) '
    'отлично потянет Qwen 3.5 35B-A3B — MoE-модель, которая активирует только 3B параметров '
    'и при этом даёт качество на уровне крупных dense-моделей.'
)

doc.add_page_break()

# ── Этап 1 ──
doc.add_heading('Этап 1. Установка Ollama', level=2)

doc.add_paragraph(
    'Ollama — самый простой способ запустить ИИ-модель локально. '
    'Работает на Linux, macOS и Windows.'
)

doc.add_heading('Установка', level=3)
add_code(doc, '# Linux / macOS:')
add_code(doc, 'curl -fsSL https://ollama.com/install.sh | sh')
add_code(doc, '')
add_code(doc, '# Windows:')
add_code(doc, '# Скачать установщик с https://ollama.com/download')
add_code(doc, '')
add_code(doc, '# Проверка:')
add_code(doc, 'ollama --version')

doc.add_heading('Первая модель', level=3)
doc.add_paragraph('Скачайте и запустите модель одной командой:')
add_code(doc, '# Рекомендация для начала — Qwen 3.5 35B-A3B (MoE, лёгкий запуск):')
add_code(doc, 'ollama pull qwen3.5:35b-a3b')
add_code(doc, '')
add_code(doc, '# Запуск диалога:')
add_code(doc, 'ollama run qwen3.5:35b-a3b')
add_code(doc, '')
add_code(doc, '# Другие актуальные модели:')
add_code(doc, 'ollama pull qwen3.5:9b        # Компактная, хороший русский')
add_code(doc, 'ollama pull qwen3.5:27b       # Dense, высокое качество')
add_code(doc, 'ollama pull deepseek-r1:14b   # Сильное рассуждение')
add_code(doc, 'ollama pull llama3.1:8b       # Meta, быстрая и лёгкая')

doc.add_heading('Проверка API', level=3)
doc.add_paragraph('Ollama запускает локальный API-сервер на порту 11434:')
add_code(doc, '# Список установленных моделей:')
add_code(doc, 'curl http://localhost:11434/api/tags')
add_code(doc, '')
add_code(doc, '# Тестовый запрос:')
add_code(doc, 'curl http://localhost:11434/api/generate -d \'{')
add_code(doc, '  "model": "qwen3.5:35b-a3b",')
add_code(doc, '  "prompt": "Что такое ПИД-регулятор?",')
add_code(doc, '  "stream": false')
add_code(doc, '}\'')

doc.add_heading('Рекомендуемые модели для инженеров', level=3)
add_table(doc,
    ['Модель', 'Размер', 'Лучшее применение', 'Команда установки'],
    [
        ['Qwen 3.5 35B-A3B', '~5 ГБ', 'Универсальная, MoE — лёгкий запуск', 'ollama pull qwen3.5:35b-a3b'],
        ['Qwen 3.5 9B', '~5.5 ГБ', 'Компактная, хороший русский', 'ollama pull qwen3.5:9b'],
        ['Qwen 3.5 27B', '~16 ГБ', 'Dense, документация, анализ', 'ollama pull qwen3.5:27b'],
        ['Qwen 3.5 122B-A10B', '~18 ГБ', 'MoE, близко к облачным', 'ollama pull qwen3.5:122b-a10b'],
        ['DeepSeek-R1 14B', '~9 ГБ', 'Сложные расчёты, рассуждения', 'ollama pull deepseek-r1:14b'],
        ['Nomic Embed Text', '~274 МБ', 'Embeddings для RAG', 'ollama pull nomic-embed-text'],
    ]
)

doc.add_page_break()

# ── Этап 2 ──
doc.add_heading('Этап 2. Интерфейсы для работы', level=2)

doc.add_heading('Вариант А: Open WebUI (веб-интерфейс)', level=3)
doc.add_paragraph(
    'Open WebUI — это веб-приложение, похожее на ChatGPT, но работающее локально.'
)
add_code(doc, '# Установка через Docker:')
add_code(doc, 'docker run -d -p 3000:8080 \\')
add_code(doc, '  -e OLLAMA_BASE_URL=http://host.docker.internal:11434 \\')
add_code(doc, '  -v open-webui:/app/backend/data \\')
add_code(doc, '  --name open-webui \\')
add_code(doc, '  --restart always \\')
add_code(doc, '  ghcr.io/open-webui/open-webui:main')
add_code(doc, '')
add_code(doc, '# Открыть в браузере: http://localhost:3000')
doc.add_paragraph()
doc.add_paragraph('Возможности Open WebUI:')
doc.add_paragraph('• Загрузка файлов (PDF, Excel, изображения) прямо в чат')
doc.add_paragraph('• История диалогов')
doc.add_paragraph('• Несколько пользователей (каждый инженер — свой аккаунт)')
doc.add_paragraph('• RAG — загрузка базы знаний для поиска по документам')
doc.add_paragraph('• Переключение между моделями в одном интерфейсе')

doc.add_heading('Вариант Б: Через OpenClaw (в Telegram)', level=3)
doc.add_paragraph(
    'Если OpenClaw уже установлен, можно подключить локальную модель как провайдер:'
)
doc.add_paragraph('1. Добавьте в конфиг OpenClaw провайдер ollama')
doc.add_paragraph('2. Укажите адрес: http://localhost:11434 (или IP вашего ПК)')
doc.add_paragraph('3. Выберите модель по умолчанию')
doc.add_paragraph(
    'Теперь ваш Telegram-бот может использовать и облачные, и локальные модели. '
    'Можно переключаться: Claude для сложных задач, локальная Qwen для рутины.'
)

doc.add_heading('Вариант В: VS Code + Continue', level=3)
doc.add_paragraph(
    'Для инженеров, работающих с кодом (ПЛК, Python, MATLAB):'
)
doc.add_paragraph('1. Установите расширение Continue в VS Code')
doc.add_paragraph('2. В настройках Continue укажите Ollama как провайдер')
doc.add_paragraph('3. Получите автодополнение кода и чат прямо в редакторе')

add_eng_example(doc,
    'Пишете программу для ПЛК на Structured Text — Continue подсказывает код, '
    'объясняет функциональные блоки и помогает отлаживать логику.'
)

doc.add_page_break()

# ── Этап 3 ──
doc.add_heading('Этап 3. Оптимизация производительности', level=2)

doc.add_heading('Квантизация — ключ к скорости', level=3)
doc.add_paragraph(
    'Квантизация — это «сжатие» модели с минимальной потерей качества. '
    'Модели бывают разной точности:'
)
add_table(doc,
    ['Квантизация', 'Размер (для 7B)', 'Качество (оценка)', 'Когда использовать'],
    [
        ['FP16 (полная)', '~14 ГБ', '100%', 'Если хватает VRAM'],
        ['Q8_0', '~7.7 ГБ', '~99%', 'Почти без потерь'],
        ['Q5_K_M', '~5.3 ГБ', '~97%', 'Хороший баланс'],
        ['Q4_K_M', '~4.4 ГБ', '~95%', 'Рекомендовано для большинства'],
        ['Q3_K_M', '~3.5 ГБ', '~90%', 'Если мало памяти'],
        ['Q2_K', '~2.8 ГБ', '~80%', 'Только если нет другого выбора'],
    ]
)
doc.add_paragraph()
add_code(doc, '# Скачать квантованную версию:')
add_code(doc, 'ollama pull qwen3.5:27b-q4_K_M')

doc.add_heading('Проверка использования GPU', level=3)
add_code(doc, '# Проверить, использует ли Ollama GPU:')
add_code(doc, 'ollama ps')
add_code(doc, '')
add_code(doc, '# Мониторинг GPU:')
add_code(doc, 'watch -n 1 nvidia-smi')

doc.add_heading('Настройки Ollama', level=3)
add_code(doc, '# Переменные окружения (добавить в ~/.bashrc или /etc/environment):')
add_code(doc, '')
add_code(doc, '# Количество параллельных запросов:')
add_code(doc, 'OLLAMA_NUM_PARALLEL=2')
add_code(doc, '')
add_code(doc, '# Размер контекста (по умолчанию 2048):')
add_code(doc, 'OLLAMA_NUM_CTX=8192')
add_code(doc, '')
add_code(doc, '# Привязка к конкретному адресу (для доступа по сети):')
add_code(doc, 'OLLAMA_HOST=0.0.0.0:11434')

doc.add_page_break()

# ── Этап 4 ──
doc.add_heading('Этап 4. Продвинутые сценарии', level=2)

doc.add_heading('RAG — поиск по своим документам', level=3)
doc.add_paragraph(
    'RAG (Retrieval-Augmented Generation) — технология, которая позволяет модели '
    'искать информацию в ваших документах перед ответом. Это как дать модели доступ '
    'к библиотеке.'
)
doc.add_paragraph('Как это работает:')
doc.add_paragraph('1. Ваши документы (PDF, DOCX, TXT) разбиваются на фрагменты')
doc.add_paragraph('2. Каждый фрагмент преобразуется в числовой вектор (embedding)')
doc.add_paragraph('3. Когда вы задаёте вопрос, система находит релевантные фрагменты')
doc.add_paragraph('4. Модель отвечает, опираясь на найденные фрагменты')

doc.add_paragraph('Настройка через Open WebUI:')
doc.add_paragraph('• Перейти в раздел «Workspace» → «Knowledge»')
doc.add_paragraph('• Создать коллекцию (например, «Каталоги КИПиА»)')
doc.add_paragraph('• Загрузить PDF-файлы')
doc.add_paragraph('• В чате выбрать эту коллекцию как источник')

add_eng_example(doc,
    'Загрузите 50 PDF-каталогов оборудования. Теперь вопрос «Какой расходомер подойдёт '
    'для серной кислоты 95%, DN50, давление 6 бар?» даст ответ с конкретными моделями '
    'из ваших каталогов, а не выдуманные позиции.'
)

doc.add_heading('Embeddings для поиска', level=3)
add_code(doc, '# Установите модель для эмбеддингов:')
add_code(doc, 'ollama pull nomic-embed-text')
add_code(doc, '')
add_code(doc, '# Она используется автоматически в Open WebUI для RAG')
add_code(doc, '# Или вручную через API:')
add_code(doc, 'curl http://localhost:11434/api/embed -d \'{')
add_code(doc, '  "model": "nomic-embed-text",')
add_code(doc, '  "input": "датчик давления для агрессивных сред"')
add_code(doc, '}\'')

doc.add_heading('Мультимодальные модели (работа с изображениями)', level=3)
doc.add_paragraph(
    'Некоторые модели понимают изображения — можно отправить фото '
    'и получить анализ:'
)
add_code(doc, '# Qwen 3.5 — нативная мультимодальная модель (vision из коробки):')
add_code(doc, 'ollama pull qwen3.5:35b-a3b   # уже поддерживает изображения')
add_code(doc, '')
add_code(doc, '# Альтернатива:')
add_code(doc, 'ollama pull llama3.2-vision:11b')
doc.add_paragraph()
doc.add_paragraph('Применение для инженера:')
doc.add_paragraph('• Распознавание текста на фото шильдиков оборудования')
doc.add_paragraph('• Анализ P&ID схем (с оговорками по точности)')
doc.add_paragraph('• Считывание показаний приборов с фото')
doc.add_paragraph('• Описание монтажных ситуаций с фотографий')

doc.add_page_break()

# ══════════════════════════════════════════════
# ЧАСТЬ 3 — ИНЖЕНЕРНОЕ ПРОЕКТИРОВАНИЕ
# ══════════════════════════════════════════════
doc.add_heading('Часть 3. ИИ в инженерном проектировании', level=1)

doc.add_heading('Сценарии применения по дисциплинам', level=2)

doc.add_heading('АСУ ТП / SCADA', level=3)
doc.add_paragraph('• Генерация описаний технологических процессов')
doc.add_paragraph('• Составление перечня входных/выходных сигналов')
doc.add_paragraph('• Написание алгоритмов управления (текстовое описание)')
doc.add_paragraph('• Генерация кода на Structured Text (IEC 61131-3)')
doc.add_paragraph('• Подготовка матрицы причин и следствий (C&E Matrix)')
doc.add_paragraph('• Описание режимов работы оборудования')
doc.add_paragraph('• Составление программы испытаний')

doc.add_heading('КИПиА (контрольно-измерительные приборы)', level=3)
doc.add_paragraph('• Подбор приборов по параметрам среды')
doc.add_paragraph('• Составление опросных листов')
doc.add_paragraph('• Расчёт сужающих устройств (диафрагм)')
doc.add_paragraph('• Подбор кабеля для КИПиА')
doc.add_paragraph('• Маркировка приборов по ГОСТ 21.208 / ГОСТ 21.408')
doc.add_paragraph('• Составление кабельного журнала')
doc.add_paragraph('• Расчёт длины импульсных трубок')

doc.add_heading('Электротехника', level=3)
doc.add_paragraph('• Расчёт токов короткого замыкания')
doc.add_paragraph('• Подбор автоматических выключателей')
doc.add_paragraph('• Расчёт сечения кабелей по допустимому нагреву')
doc.add_paragraph('• Составление однолинейных схем (текстовое описание)')
doc.add_paragraph('• Расчёт компенсации реактивной мощности')
doc.add_paragraph('• Проверка селективности защит')
doc.add_paragraph('• Составление спецификации электрооборудования')

doc.add_heading('Общепроектные задачи', level=3)
doc.add_paragraph('• Написание разделов пояснительных записок')
doc.add_paragraph('• Составление ТЗ на разработку')
doc.add_paragraph('• Подготовка технико-коммерческих предложений (ТКП)')
doc.add_paragraph('• Перевод технической документации (EN↔RU)')
doc.add_paragraph('• Составление писем подрядчикам и заказчикам')
doc.add_paragraph('• Рецензирование и вычитка документов')
doc.add_paragraph('• Формирование сводных таблиц и отчётов из данных')

doc.add_page_break()

doc.add_heading('Примеры промптов для инженерных задач', level=2)

prompts = [
    ('Подбор оборудования',
     'Подбери датчик температуры для измерения температуры пара. '
     'Условия: температура до 250°C, давление 16 бар, среда — перегретый пар, '
     'присоединение — фланцевое DN25 PN40, выходной сигнал 4-20 мА + HART. '
     'Предложи 2-3 варианта от разных производителей с обоснованием.'),

    ('Генерация кода ПЛК',
     'Напиши функциональный блок на Structured Text (IEC 61131-3) для ПИД-регулятора '
     'уровня в ёмкости. Входы: PV (текущий уровень, REAL), SP (уставка, REAL). '
     'Выход: OUT (управляющий сигнал 0-100%, REAL). Параметры Kp, Ki, Kd настраиваемые. '
     'Добавь защиту от integral windup.'),

    ('Пояснительная записка',
     'Напиши раздел пояснительной записки «Система автоматизации» для проекта '
     'модернизации насосной станции. Включи: описание объекта автоматизации, '
     'перечень контролируемых параметров, структуру АСУ ТП (3 уровня), '
     'описание функций системы, требования к надёжности. Стиль — ГОСТ 34.'),

    ('Расчёт кабеля',
     'Рассчитай сечение кабеля для питания электродвигателя насоса. '
     'Данные: мощность 75 кВт, напряжение 380 В, cos φ = 0.85, '
     'длина трассы 120 м, прокладка в лотке, 4 кабеля в лотке, '
     'температура окружающей среды 35°C. Марка кабеля — ВВГнг(А)-LS. '
     'Показать расчёт по нагреву и проверку по потере напряжения.'),

    ('Матрица причин и следствий',
     'Составь матрицу причин и следствий (C&E Matrix) для ёмкости '
     'хранения нефтепродуктов. Контролируемые параметры: уровень (HH, H, L, LL), '
     'температура (HH, H), давление (HH, H), загазованность (1 и 2 порог). '
     'Указать действия: сигнализация, блокировка, закрытие клапана, '
     'остановка насоса, запуск вентиляции.'),
]

for title, prompt in prompts:
    p = doc.add_paragraph()
    run = p.add_run(f'📋 {title}')
    run.bold = True
    run.font.size = Pt(11)
    add_code(doc, prompt)
    doc.add_paragraph()

doc.add_page_break()

doc.add_heading('Сравнение моделей для инженерных задач', level=2)

add_table(doc,
    ['Задача', 'Лучшая облачная', 'Лучшая локальная', 'Комментарий'],
    [
        ['Технические тексты (РУС)', 'Claude Sonnet 4', 'Qwen 3.5 27B', 'Qwen хорошо знает русский (оценка автора)'],
        ['Код ПЛК (ST/SCL)', 'Claude Opus 4 / GPT-4.1', 'Qwen3-Coder (через OpenRouter)', 'Облачные значительно лучше (оценка автора)'],
        ['Расчёты с формулами', 'o3 / Claude Opus 4', 'DeepSeek-R1 32B', 'DeepSeek хорош в рассуждениях'],
        ['Подбор оборудования', 'Claude + RAG', 'Qwen 3.5 27B + RAG', 'Качество зависит от базы знаний'],
        ['Перевод тех. текста', 'Claude Sonnet 4', 'Qwen 3.5 27B', 'Локальная справляется хорошо (оценка автора)'],
        ['Анализ изображений', 'GPT-4.1 Vision', 'Qwen 3.5 35B-A3B (vision)', 'Qwen 3.5 — нативный vision'],
        ['Генерация документов', 'Claude Opus 4', 'Qwen 3.5 122B-A10B', 'Для ГОСТ-формата нужна настройка'],
    ]
)

doc.add_paragraph()
add_note(doc,
    'Для критически важных задач (расчёты, безопасность) всегда проверяйте результаты ИИ. '
    'Модель — это инструмент для ускорения, а не замена инженерной экспертизы. '
    'Ответственность за проектные решения несёт инженер.'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# ПРИЛОЖЕНИЯ
# ══════════════════════════════════════════════
doc.add_heading('Приложения', level=1)

doc.add_heading('Приложение А. Глоссарий', level=2)
glossary = [
    ('API (Application Programming Interface)', 'Программный интерфейс для взаимодействия с сервисом. API-ключ — «пароль» для доступа к облачной модели.'),
    ('Embedding', 'Числовое представление текста в виде вектора. Используется для поиска похожих фрагментов (RAG).'),
    ('GPU (Graphics Processing Unit)', 'Видеокарта. Ускоряет работу нейросетей в 10–50 раз по сравнению с CPU.'),
    ('LLM (Large Language Model)', 'Большая языковая модель — нейросеть, обученная на текстах, способная генерировать текст.'),
    ('Ollama', 'Open-source инструмент для запуска LLM локально. Упрощает скачивание и использование моделей.'),
    ('OpenClaw', 'Open-source платформа для создания персонального ИИ-ассистента с доступом к инструментам.'),
    ('Промпт (Prompt)', 'Текстовый запрос к ИИ-модели. Качество промпта определяет качество ответа.'),
    ('Квантизация', 'Сжатие модели путём уменьшения точности чисел (FP16 → INT4). Уменьшает размер и ускоряет работу.'),
    ('RAG (Retrieval-Augmented Generation)', 'Технология дополнения ответов модели информацией из ваших документов.'),
    ('Токен', 'Единица текста для модели. ~4 символа на английском, ~2–3 символа на русском.'),
    ('VRAM', 'Видеопамять GPU. Определяет максимальный размер модели, которую можно загрузить.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    run = p.add_run(f'{term} — ')
    run.bold = True
    p.add_run(definition)

doc.add_heading('Приложение Б. Полезные ссылки', level=2)
links = [
    ('OpenClaw', 'https://github.com/openclaw/openclaw', 'Платформа для ИИ-ассистента'),
    ('Ollama', 'https://ollama.com', 'Локальный запуск моделей'),
    ('Open WebUI', 'https://github.com/open-webui/open-webui', 'Веб-интерфейс для Ollama'),
    ('Hugging Face', 'https://huggingface.co', 'Каталог моделей и датасетов'),
    ('Anthropic', 'https://console.anthropic.com', 'API-ключи для Claude'),
    ('OpenRouter', 'https://openrouter.ai', 'Единый доступ ко всем моделям'),
    ('Continue', 'https://continue.dev', 'ИИ-ассистент для VS Code'),
    ('ClawHub', 'https://clawhub.com', 'Навыки для OpenClaw'),
]
for name, url, desc in links:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}')
    run.bold = True
    p.add_run(f' — {url} — {desc}')

doc.add_heading('Приложение В. Чек-лист безопасности при работе с ИИ', level=2)
security_items = [
    'Не отправляйте в облачные модели чертежи с грифом «коммерческая тайна» — используйте локальную модель',
    'API-ключи храните в переменных окружения, не в коде',
    'Для конфиденциальных проектов — только локальная модель или приватный API',
    'Всегда проверяйте расчёты ИИ вручную перед использованием в проекте',
    'Не используйте ИИ для генерации финальных чертежей — только черновики и наброски',
    'Указывайте в документации, что ИИ использовался как вспомогательный инструмент',
    'Регулярно обновляйте ПО (Ollama, OpenClaw) для закрытия уязвимостей',
    'Настройте firewall на сервере с Ollama, если он доступен по сети',
]
for item in security_items:
    doc.add_paragraph(f'☐ {item}')

doc.add_heading('Приложение Г. Типовой бюджет', level=2)
add_table(doc,
    ['Статья', 'Вариант «Минимум»', 'Вариант «Оптимум»', 'Вариант «Максимум»'],
    [
        ['VPS для OpenClaw', '300 ₽/мес', '500 ₽/мес', '1000 ₽/мес'],
        ['API облачной модели', 'Бесплатно (OpenRouter free)', '1500 ₽/мес', '5000 ₽/мес'],
        ['Локальная модель', 'Бесплатно', 'Бесплатно', 'Бесплатно'],
        ['Железо для лок. модели', 'Есть ПК', 'Есть ПК + GPU', 'Выделенный GPU-сервер'],
        ['ИТОГО (ежемесячно)', '~300 ₽', '~2000 ₽', '~6000+ ₽'],
    ]
)
doc.add_paragraph()
add_note(doc,
    'Для сравнения: лицензия на один коммерческий ИИ-ассистент (ChatGPT Plus, Claude Pro) '
    'стоит $20/мес (~1800 ₽). OpenClaw + бесплатные модели через OpenRouter '
    '(Qwen3-Coder 480B, DeepSeek R1) = полностью бесплатный вариант (кроме VPS). '
    'Цены приблизительные, актуальны на март 2026.'
)

# ── Save ──
output_path = '/root/.openclaw/workspace/agents/main/docs/AI_Implementation_Guide_for_Engineers.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
print(f'Size: {os.path.getsize(output_path)} bytes')
