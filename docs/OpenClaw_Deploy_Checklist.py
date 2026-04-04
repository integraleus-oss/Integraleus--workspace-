#!/usr/bin/env python3
"""Generate OpenClaw Deploy Checklist as Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_checkbox(doc, text, bold_prefix=None, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('☐ ')
    run.font.size = Pt(11)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(11)
        rt = p.add_run(text)
        rt.font.size = Pt(11)
    else:
        rt = p.add_run(text)
        rt.font.size = Pt(11)
    return p

def add_code(doc, code, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent + 0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_note(doc, text, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

# ============================================
# TITLE
# ============================================
title = doc.add_heading('Чеклист: развёртывание ИИ-агента на OpenClaw', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Пошаговая инструкция для VPS и локального сервера')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Версия 1.0 • Март 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ============================================
# PART 1
# ============================================
doc.add_heading('Часть 1. Подготовка сервера', level=1)

doc.add_heading('1.1 Минимальные требования', level=2)
make_table(doc,
    ['Параметр', 'VPS', 'Локальный сервер'],
    [
        ['ОС', 'Ubuntu 22.04+ / Debian 12+', 'Любой Linux с systemd (вкл. Astra, РЕД ОС, ALT)'],
        ['RAM', '≥ 2 GB', '≥ 2 GB'],
        ['Диск', '≥ 10 GB', '≥ 10 GB'],
        ['CPU', '1–2 ядра', '1–2 ядра'],
        ['Сеть', 'Публичный IP, доступ к API LLM', 'Доступ наружу к API LLM-провайдеров'],
        ['Node.js', 'v22+', 'v22+'],
    ]
)

doc.add_heading('1.2 Установка ОС и базовых пакетов', level=2)
add_checkbox(doc, 'Обновить систему')
add_code(doc, 'apt update && apt upgrade -y')
add_checkbox(doc, 'Установить базовые пакеты')
add_code(doc, 'apt install -y curl git build-essential')
add_checkbox(doc, 'Установить Node.js 22+')
add_code(doc, 'curl -fsSL https://deb.nodesource.com/setup_22.x | bash -\napt install -y nodejs\nnode --version  # должно быть v22+')

doc.add_heading('1.3 Безопасность сервера', level=2)
add_checkbox(doc, 'SSH только по ключам')
add_code(doc, '# /etc/ssh/sshd_config\nPasswordAuthentication no\nPermitRootLogin without-password')
add_checkbox(doc, 'Настроить фаервол')
add_code(doc, 'ufw allow 22/tcp\nufw allow 443/tcp\nufw enable')
add_checkbox(doc, 'Установить fail2ban')
add_code(doc, 'apt install -y fail2ban && systemctl enable fail2ban')
add_checkbox(doc, 'Создать отдельного пользователя (рекомендуется)')
add_code(doc, 'useradd -m -s /bin/bash openclaw')
add_checkbox(doc, 'Автообновления безопасности')
add_code(doc, 'apt install -y unattended-upgrades')

# ============================================
# PART 2
# ============================================
doc.add_heading('Часть 2. Установка OpenClaw', level=1)
add_checkbox(doc, 'Установить OpenClaw')
add_code(doc, 'npm install -g openclaw')
add_checkbox(doc, 'Проверить версию')
add_code(doc, 'openclaw --version')
add_checkbox(doc, 'Запустить визард настройки')
add_code(doc, 'openclaw configure')

# ============================================
# PART 3
# ============================================
doc.add_heading('Часть 3. Настройка LLM-провайдера', level=1)
add_checkbox(doc, 'Выбрать провайдера и получить API-ключ')

make_table(doc,
    ['Провайдер', 'Модели', 'Примерная стоимость'],
    [
        ['Anthropic', 'Claude Sonnet / Opus', '$3–15 / 1M токенов'],
        ['OpenAI', 'GPT-5 / o3', '$2–15 / 1M токенов'],
        ['OpenRouter', 'Любые (прокси)', 'Зависит от модели'],
    ]
)
doc.add_paragraph()

add_checkbox(doc, 'Настроить провайдер: ', 'openclaw configure → ', 0.5)
add_checkbox(doc, 'Выбрать модель по умолчанию (рекомендация: Sonnet для баланса цена/качество)')
add_note(doc, '💰 При активном использовании (50–100 сообщений/день) ≈ $30–100/мес на Sonnet')

# ============================================
# PART 4
# ============================================
doc.add_heading('Часть 4. Настройка доступа', level=1)

doc.add_heading('Вариант A: Telegram-бот', level=2)
add_note(doc, 'Проще всего для быстрого старта.')
add_checkbox(doc, 'Создать бота: @BotFather → /newbot → получить токен')
add_checkbox(doc, 'Добавить канал: openclaw configure → секция каналов → Telegram')
add_checkbox(doc, 'Вставить токен бота')
add_checkbox(doc, 'Указать allowedUsers — Telegram ID авторизованных пользователей')
add_checkbox(doc, 'Тест: написать боту → должен ответить')

doc.add_heading('Вариант B: Веб-интерфейс (Control UI)', level=2)
add_note(doc, 'Для корпоративного использования через браузер.')

add_checkbox(doc, 'Настроить gateway.bind:')
add_note(doc, 'VPS: "auto" или "lan" • Локальный: "lan" (LAN) или "loopback" (только localhost)')

add_checkbox(doc, 'Настроить аутентификацию (gateway.auth):')
add_note(doc, 'mode: "token" — по токену | "password" — логин/пароль | "trusted-proxy" — за корп. SSO')

add_checkbox(doc, 'Включить Control UI: gateway.controlUi.enabled: true')

doc.add_heading('Веб-доступ через интернет (VPS)', level=3)
add_checkbox(doc, 'Настроить домен (например agent.company.ru)')
add_checkbox(doc, 'SSL — Let\'s Encrypt:')
add_code(doc, 'apt install -y certbot python3-certbot-nginx\ncertbot --nginx -d agent.company.ru')
add_checkbox(doc, 'Reverse proxy (Nginx):')
add_code(doc, '''server {
    listen 443 ssl;
    server_name agent.company.ru;

    ssl_certificate     /etc/letsencrypt/live/agent.company.ru/fullchain.pem;
    ssl_certificate_key /etc/letsencrypt/live/agent.company.ru/privkey.pem;

    location / {
        proxy_pass http://127.0.0.1:19000;
        proxy_http_version 1.1;
        proxy_set_header Upgrade $http_upgrade;
        proxy_set_header Connection "upgrade";
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_read_timeout 86400;
    }
}''')
add_checkbox(doc, 'Закрыть прямой доступ к порту gateway')
add_code(doc, 'ufw deny 19000')

doc.add_heading('Веб-доступ из LAN (локальный сервер)', level=3)
add_checkbox(doc, 'SSL необязателен (но рекомендуется самоподписанный)')
add_checkbox(doc, 'Nginx необязателен — доступ напрямую http://192.168.x.x:19000')
add_checkbox(doc, 'Достаточно gateway.auth.mode: "token" или "password"')

doc.add_heading('Вариант C: Telegram + Web одновременно', level=2)
add_note(doc, 'OpenClaw поддерживает несколько каналов параллельно — можно включить оба.')

# ============================================
# PART 5
# ============================================
doc.add_heading('Часть 5. Настройка агента (личность и знания)', level=1)
add_checkbox(doc, 'Перейти в workspace: ~/.openclaw/workspace/agents/main/')

make_table(doc,
    ['Файл', 'Что писать', 'Обязательно?'],
    [
        ['SOUL.md', 'Личность, тон, правила поведения агента', '✅ Да'],
        ['AGENTS.md', 'Системные правила: что можно/нельзя', '✅ Да'],
        ['USER.md', 'Профиль пользователя(ей)', 'Рекомендуется'],
        ['TOOLS.md', 'Заметки по инфраструктуре, сервисам', 'По ситуации'],
        ['HEARTBEAT.md', 'Периодические проверки, мониторинг', 'По ситуации'],
        ['docs/', 'База знаний — документы, инструкции', 'По ситуации'],
    ]
)
doc.add_paragraph()

add_checkbox(doc, 'Установить нужные скиллы:')
add_code(doc, 'clawhub search <тема>\nclawhub install <skill-name>')

# ============================================
# PART 6
# ============================================
doc.add_heading('Часть 6. Политика безопасности агента', level=1)
add_checkbox(doc, ' — что агент может выполнять на сервере', 'Exec-политика')
add_note(doc, 'ask — спрашивает перед каждой командой | allowlist — список разрешённых | full — всё (⚠️ только для доверенных сред)')
add_checkbox(doc, ' — список авторизованных отправителей', 'allowedUsers')
add_checkbox(doc, ' — хранятся в конфиге OpenClaw, НЕ в workspace-файлах', 'API-ключи')
add_checkbox(doc, ' — ограничение доступных инструментов на уровне gateway', 'Gateway tools policy')

# ============================================
# PART 7
# ============================================
doc.add_heading('Часть 7. Автозапуск (systemd)', level=1)
add_checkbox(doc, 'Создать юнит-файл:')
add_code(doc, '''# /etc/systemd/system/openclaw-gateway.service
[Unit]
Description=OpenClaw Gateway
After=network-online.target
Wants=network-online.target

[Service]
Type=simple
User=openclaw
ExecStart=/usr/bin/openclaw gateway start --foreground
Restart=always
RestartSec=10
Environment=NODE_ENV=production
WorkingDirectory=/home/openclaw

[Install]
WantedBy=multi-user.target''')

add_checkbox(doc, 'Активировать:')
add_code(doc, 'systemctl daemon-reload\nsystemctl enable --now openclaw-gateway')
add_checkbox(doc, 'Проверить статус')
add_code(doc, 'systemctl status openclaw-gateway')

# ============================================
# PART 8
# ============================================
doc.add_heading('Часть 8. Проверка и тестирование', level=1)
add_checkbox(doc, 'openclaw status — все компоненты зелёные')
add_checkbox(doc, 'Отправить тестовое сообщение (Telegram / веб)')
add_checkbox(doc, 'Агент ответил осмысленно')
add_checkbox(doc, 'Проверить exec: попросить агента выполнить uname -a')
add_checkbox(doc, 'Проверить heartbeat (если настроен)')
add_checkbox(doc, 'Kill-тест: убить процесс → systemd перезапустил')
add_checkbox(doc, 'Reboot-тест: перезагрузить сервер → gateway поднялся сам')

# ============================================
# PART 9
# ============================================
doc.add_heading('Часть 9. Эксплуатация', level=1)
add_checkbox(doc, ' — openclaw backup по расписанию', 'Бэкапы:')
add_checkbox(doc, ' — openclaw update (ручной или по расписанию)', 'Обновления:')
add_checkbox(doc, ' — journalctl -u openclaw-gateway -f', 'Логи:')
add_checkbox(doc, ' — настроить heartbeat-проверки в HEARTBEAT.md', 'Мониторинг:')

# ============================================
# SUMMARY TABLE
# ============================================
doc.add_heading('Сводка различий: VPS vs Локальный сервер', level=1)
make_table(doc,
    ['Аспект', 'VPS', 'Локальный сервер'],
    [
        ['IP', 'Публичный, статический', 'Внутренний (192.168.x.x)'],
        ['Домен', 'Нужен для веб-доступа', 'Не обязателен'],
        ['SSL', 'Обязателен (Let\'s Encrypt)', 'Рекомендуется (самоподписанный)'],
        ['Nginx', 'Нужен как reverse proxy', 'Не обязателен'],
        ['Фаервол', 'Критичен (публичный IP)', 'Важен, но менее критичен'],
        ['Telegram', 'Работает напрямую', 'Нужен доступ к api.telegram.org'],
        ['Веб-доступ', 'Через домен + SSL', 'Через LAN IP напрямую'],
        ['Корп. прокси', 'Не актуально', 'Может потребоваться HTTP_PROXY'],
        ['Ответственность', 'Хостер + вы', 'Полностью вы'],
    ]
)

# ============================================
# SAVE
# ============================================
out = '/root/.openclaw/workspace/agents/main/docs/OpenClaw_Deploy_Checklist.docx'
doc.save(out)
print(f'Saved: {out}')
